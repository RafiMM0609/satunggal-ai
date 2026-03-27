"""
DocxEditorTool – Edit dokumen .docx dengan presisi XML-level.

Kemampuan:
  1. replace_text      – Ganti teks spesifik di dalam Run tanpa menyentuh format
  2. add_paragraph     – Tambah paragraf baru dengan mewarisi style paragraf sebelumnya
  3. delete_paragraph  – Hapus paragraf berdasarkan indeks (0-based)
  4. replace_paragraph – Ganti konten paragraf sambil mempertahankan paragraph properties

Prinsip kerja (sesuai brief):
  - Parsing XML: membaca struktur internal word/document.xml melalui python-docx
  - Object Mapping: Paragraph → Run → Text sebagai hirarki DOM
  - Style Extraction: deep-copy <w:pPr> dan <w:rPr> dari paragraf referensi
  - Precise Injection: hanya mengubah <w:t> di dalam <w:r> yang tepat
  - Inheritance: paragraf baru mewarisi numId, ilvl, dan StyleID dari referensi
  - Re-packaging: simpan kembali menjadi file .docx yang valid

Input (via task.metadata):
  docx_path (str):       Path ke file .docx original
  docx_edits (list):     Daftar operasi edit. Setiap item adalah dict:
      {"op": "replace_text",  "find": str, "replace": str,
       "paragraph_index": int|None}        # None = seluruh dokumen
      {"op": "add_paragraph", "text": str, "after_paragraph_index": int,
       "style_from_index": int|None}       # None = inherit dari after_paragraph_index
      {"op": "delete_paragraph", "paragraph_index": int}
      {"op": "replace_paragraph", "paragraph_index": int, "new_text": str}
  output_docx_path (str, optional): Path untuk menyimpan file hasil edit.
      Default: menimpa file asli dengan suffix "_edited".

Output dict:
  {
      "edited_docx_path": str,
      "changes_made": int,
      "details": list[str],
  }
  Atau {"error": str} jika gagal.
"""

from __future__ import annotations

import copy
import logging
import os
import tempfile
from datetime import datetime
from typing import Any, TYPE_CHECKING

from src.tools.base_tool import BaseTool

if TYPE_CHECKING:
    from src.memory.state import AgentTask

logger = logging.getLogger(__name__)

# XML namespace untuk Word Processing ML
_W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
_XML_NS = "http://www.w3.org/XML/1998/namespace"


def _qn(tag: str) -> str:
    """Buat qualified name dari tag Word (misal 'w:pPr' → '{...}pPr')."""
    prefix, local = tag.split(":", 1)
    if prefix == "w":
        return f"{{{_W_NS}}}{local}"
    if prefix == "xml":
        return f"{{{_XML_NS}}}{local}"
    return tag


def _make_output_path(docx_path: str, session_id: str) -> str:
    """Buat path output untuk file .docx hasil edit."""
    out_dir = os.path.join(tempfile.gettempdir(), "advance_ai_docx_edited", session_id)
    os.makedirs(out_dir, exist_ok=True)
    basename = os.path.basename(docx_path)
    name, ext = os.path.splitext(basename)
    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    return os.path.join(out_dir, f"{name}_edited_{ts}{ext}")


# ── Core editing functions ────────────────────────────────────────────────────

def _replace_text_cross_runs(para, find: str, replace: str, actual_idx: int) -> list[str]:
    """
    Ganti teks yang tersebar di beberapa Run dalam satu paragraf.

    Strategi:
    - Bangun peta karakter → run untuk seluruh teks paragraf.
    - Temukan posisi `find` dalam teks lengkap.
    - Tentukan run mana saja yang terlibat.
    - Simpan teks pengganti ke run PERTAMA yang terlibat.
    - Kosongkan teks di run-run berikutnya yang merupakan bagian dari `find`.

    Returns:
        list[str]: Deskripsi perubahan yang dilakukan.
    """
    full_text = para.text
    find_start = full_text.find(find)
    if find_start == -1:
        return []

    find_end = find_start + len(find)

    # Bangun rentang karakter per run: [(run_start, run_end, run_index), ...]
    runs = para.runs
    run_ranges: list[tuple[int, int, int]] = []
    pos = 0
    for run_idx, run in enumerate(runs):
        run_len = len(run.text)
        run_ranges.append((pos, pos + run_len, run_idx))
        pos += run_len

    # Tentukan run mana yang beririsan dengan posisi `find`
    involved: list[int] = [
        run_idx
        for (r_start, r_end, run_idx) in run_ranges
        if r_start < find_end and r_end > find_start
    ]
    if not involved:
        return []

    first_run_idx = involved[0]
    last_run_idx  = involved[-1]

    first_run_char_start = run_ranges[first_run_idx][0]
    last_run_char_end    = run_ranges[last_run_idx][1]

    # Teks sebelum `find` dalam run pertama
    prefix_in_first = full_text[first_run_char_start:find_start]
    # Teks setelah `find` dalam run terakhir
    suffix_in_last  = full_text[find_end:last_run_char_end]

    new_first_run_text = prefix_in_first + replace + suffix_in_last
    old_combined_text  = full_text[first_run_char_start:last_run_char_end]

    # Terapkan perubahan
    runs[first_run_idx].text = new_first_run_text
    for run_idx in involved[1:]:
        runs[run_idx].text = ""

    return [
        f"[¶{actual_idx}] Run(cross): '{old_combined_text}' → '{new_first_run_text}'"
    ]


def _replace_text_in_run_level(
    paragraphs: list,
    find: str,
    replace: str,
    paragraph_index: int | None,
) -> list[str]:
    """
    Ganti teks di level Run (<w:t>) tanpa menyentuh <w:rPr> (format).

    Langkah 1: Cari `find` di dalam setiap run secara individual (aman untuk
               formatting bold/italic per-run).
    Langkah 2: Jika tidak ditemukan di run manapun tetapi ada di para.text
               (artinya teks terbagi antar-run), gunakan _replace_text_cross_runs
               sebagai fallback.

    Returns:
        list[str]: Deskripsi perubahan yang dilakukan.
    """
    details: list[str] = []
    targets = [paragraphs[paragraph_index]] if paragraph_index is not None else paragraphs

    for para_idx, para in enumerate(targets):
        actual_idx = paragraph_index if paragraph_index is not None else para_idx

        # Langkah 1: coba per-run (preserves per-run formatting)
        replaced_in_single_run = False
        for run in para.runs:
            if find in run.text:
                old_text = run.text
                run.text = run.text.replace(find, replace)
                details.append(
                    f"[¶{actual_idx}] Run: '{old_text}' → '{run.text}'"
                )
                replaced_in_single_run = True

        # Langkah 2: fallback cross-run jika tidak ditemukan di run tunggal
        if not replaced_in_single_run and find in para.text:
            cross_details = _replace_text_cross_runs(para, find, replace, actual_idx)
            details.extend(cross_details)

    return details


def _add_paragraph_after(
    doc,
    after_paragraph_index: int,
    text: str,
    style_from_index: int | None,
) -> list[str]:
    """
    Tambah paragraf baru setelah indeks tertentu, mewarisi style dari referensi.

    Prinsip Inheritance:
    - Deep-copy <w:pPr> dari paragraf referensi (termasuk numId, ilvl, pStyle)
    - Deep-copy <w:rPr> dari run pertama paragraf referensi
    - Word akan melanjutkan urutan numbering secara otomatis karena numId sama

    Returns:
        list[str]: Deskripsi perubahan yang dilakukan.
    """
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    paragraphs = doc.paragraphs
    if after_paragraph_index >= len(paragraphs):
        return [f"[ERROR] after_paragraph_index {after_paragraph_index} di luar jangkauan"]

    ref_para = paragraphs[after_paragraph_index]
    style_para_idx = style_from_index if style_from_index is not None else after_paragraph_index
    style_para_idx = max(0, min(style_para_idx, len(paragraphs) - 1))
    style_para = paragraphs[style_para_idx]

    # Buat elemen paragraf baru
    new_p = OxmlElement("w:p")

    # ── Style Extraction + Inheritance: Deep-copy <w:pPr> ─────────────────
    pPr = style_para._element.find(qn("w:pPr"))
    if pPr is not None:
        new_pPr = copy.deepcopy(pPr)
        new_p.append(new_pPr)

    # ── Buat Run dengan teks ───────────────────────────────────────────────
    new_r = OxmlElement("w:r")

    # Warisi <w:rPr> dari run pertama paragraf referensi
    if style_para.runs:
        rPr = style_para.runs[0]._element.find(qn("w:rPr"))
        if rPr is not None:
            new_r.append(copy.deepcopy(rPr))

    # Tambah <w:t> dengan teks
    new_t = OxmlElement("w:t")
    new_t.text = text
    # Pertahankan spasi di awal/akhir
    if text != text.strip():
        new_t.set(f"{{{_XML_NS}}}space", "preserve")
    new_r.append(new_t)
    new_p.append(new_r)

    # ── Precise Injection: Sisipkan setelah paragraf referensi ────────────
    ref_para._element.addnext(new_p)

    return [
        f"[¶{after_paragraph_index}+] Tambah paragraf baru: '{text[:60]}'"
        f"{'...' if len(text) > 60 else ''} "
        f"(inherit style dari ¶{style_para_idx})"
    ]


def _delete_paragraph(doc, paragraph_index: int) -> list[str]:
    """
    Hapus paragraf berdasarkan indeks (0-based).

    Returns:
        list[str]: Deskripsi perubahan yang dilakukan.
    """
    paragraphs = doc.paragraphs
    if paragraph_index >= len(paragraphs):
        return [f"[ERROR] paragraph_index {paragraph_index} di luar jangkauan"]

    para = paragraphs[paragraph_index]
    text_preview = para.text[:60]
    parent = para._element.getparent()
    if parent is None:
        return [f"[ERROR] Paragraf ¶{paragraph_index} tidak memiliki parent"]

    parent.remove(para._element)
    return [f"[¶{paragraph_index}] Hapus paragraf: '{text_preview}'"]


def _replace_paragraph_content(
    doc, paragraph_index: int, new_text: str
) -> list[str]:
    """
    Ganti seluruh konten paragraf sambil mempertahankan <w:pPr> (format paragraf).

    Prinsip Precise Editing:
    - Hapus semua <w:r> (run) yang lama
    - Pertahankan <w:pPr> (paragraph properties: style, numbering, indentasi)
    - Buat run baru dengan <w:rPr> dari run pertama lama (jika ada)

    Returns:
        list[str]: Deskripsi perubahan yang dilakukan.
    """
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    paragraphs = doc.paragraphs
    if paragraph_index >= len(paragraphs):
        return [f"[ERROR] paragraph_index {paragraph_index} di luar jangkauan"]

    para = paragraphs[paragraph_index]
    old_text_preview = para.text[:60]

    # ── Style Extraction: Simpan <w:rPr> dari run pertama ─────────────────
    first_rPr = None
    if para.runs:
        rPr = para.runs[0]._element.find(qn("w:rPr"))
        if rPr is not None:
            first_rPr = copy.deepcopy(rPr)

    # ── Hapus semua run lama (pertahankan <w:pPr>) ────────────────────────
    for child in list(para._element):
        if child.tag != qn("w:pPr"):
            para._element.remove(child)

    # ── Buat run baru dengan teks ──────────────────────────────────────────
    new_r = OxmlElement("w:r")
    if first_rPr is not None:
        new_r.append(first_rPr)

    new_t = OxmlElement("w:t")
    new_t.text = new_text
    if new_text != new_text.strip():
        new_t.set(f"{{{_XML_NS}}}space", "preserve")
    new_r.append(new_t)
    para._element.append(new_r)

    return [
        f"[¶{paragraph_index}] Ganti konten: "
        f"'{old_text_preview}' → '{new_text[:60]}'"
    ]


# ── Main editing function ─────────────────────────────────────────────────────

def apply_docx_edits(
    docx_path: str,
    edits: list[dict],
    output_path: str,
    session_id: str = "default",
) -> dict[str, Any]:
    """
    Terapkan daftar operasi edit pada file .docx dengan presisi XML-level.

    Args:
        docx_path:   Path ke file .docx asli.
        edits:       Daftar operasi edit.
        output_path: Path untuk file hasil edit.
        session_id:  Identifier sesi (untuk logging).

    Returns:
        Dict berisi edited_docx_path, changes_made, details.
    """
    try:
        from docx import Document
    except ImportError as exc:
        raise ImportError("python-docx tidak terinstal.") from exc

    if not os.path.isfile(docx_path):
        return {"error": f"File tidak ditemukan: {docx_path}"}

    doc = Document(docx_path)
    all_details: list[str] = []
    changes_made = 0

    for i, edit in enumerate(edits):
        op = edit.get("op", "")
        try:
            if op == "replace_text":
                details = _replace_text_in_run_level(
                    paragraphs=doc.paragraphs,
                    find=edit["find"],
                    replace=edit["replace"],
                    paragraph_index=edit.get("paragraph_index"),
                )
                all_details.extend(details)
                changes_made += len(details)

            elif op == "add_paragraph":
                details = _add_paragraph_after(
                    doc=doc,
                    after_paragraph_index=edit["after_paragraph_index"],
                    text=edit["text"],
                    style_from_index=edit.get("style_from_index"),
                )
                all_details.extend(details)
                if not any(d.startswith("[ERROR]") for d in details):
                    changes_made += 1

            elif op == "delete_paragraph":
                details = _delete_paragraph(
                    doc=doc,
                    paragraph_index=edit["paragraph_index"],
                )
                all_details.extend(details)
                if not any(d.startswith("[ERROR]") for d in details):
                    changes_made += 1

            elif op == "replace_paragraph":
                details = _replace_paragraph_content(
                    doc=doc,
                    paragraph_index=edit["paragraph_index"],
                    new_text=edit["new_text"],
                )
                all_details.extend(details)
                if not any(d.startswith("[ERROR]") for d in details):
                    changes_made += 1

            else:
                all_details.append(f"[Edit {i}] Operasi tidak dikenal: '{op}'")

        except (KeyError, IndexError) as exc:
            all_details.append(f"[Edit {i}] Error pada operasi '{op}': {exc}")
            logger.warning(
                "docx_editor: edit #%d op=%r error=%s session=%s",
                i, op, exc, session_id,
            )

    # ── Re-packaging: simpan kembali menjadi .docx ────────────────────────
    doc.save(output_path)
    logger.info(
        "docx_editor: saved edited file to %s | changes=%d session=%s",
        output_path, changes_made, session_id,
    )

    return {
        "edited_docx_path": output_path,
        "changes_made": changes_made,
        "details": all_details,
    }


def get_paragraph_map(docx_path: str) -> list[dict[str, Any]]:
    """
    Ambil daftar semua paragraf dari .docx dengan indeks (0-based).

    Returns:
        list[dict] dengan kunci: index, text, style, is_heading, is_list
    """
    try:
        from docx import Document
    except ImportError as exc:
        raise ImportError("python-docx tidak terinstal.") from exc

    doc = Document(docx_path)
    result = []
    for i, para in enumerate(doc.paragraphs):
        style_name = para.style.name if para.style else "Normal"
        is_heading = style_name.lower().startswith("heading")
        is_list = para._element.find(f".//{_qn('w:numPr')}") is not None
        result.append({
            "index": i,
            "text": para.text,
            "style": style_name,
            "is_heading": is_heading,
            "is_list": is_list,
        })
    return result


# ── Tool class ────────────────────────────────────────────────────────────────

class DocxEditorTool(BaseTool):
    """
    Tool untuk mengedit file .docx dengan presisi XML-level.

    Input (via task.metadata):
        docx_path (str):        Path ke file .docx asli.
        docx_edits (list):      Daftar operasi edit (dict per operasi).
        output_docx_path (str, optional): Path output. Default: auto-generated.

    Output:
        {
            "edited_docx_path": str,
            "changes_made": int,
            "details": list[str],
        }
    """

    name = "docx_editor"

    async def run(self, task: "AgentTask") -> dict[str, Any]:
        docx_path: str | None = task.metadata.get("docx_path")
        edits: list | None = task.metadata.get("docx_edits")

        if not docx_path:
            return {"error": "docx_path tidak ditemukan di task.metadata"}
        if edits is None:
            return {"error": "docx_edits tidak ditemukan di task.metadata"}

        output_path: str = task.metadata.get(
            "output_docx_path",
            _make_output_path(docx_path, task.session_id),
        )

        try:
            result = apply_docx_edits(
                docx_path=docx_path,
                edits=edits,
                output_path=output_path,
                session_id=task.session_id,
            )
            if "edited_docx_path" in result:
                task.metadata["document_path"] = result["edited_docx_path"]
            return result
        except FileNotFoundError as exc:
            logger.error("DocxEditorTool: %s", exc)
            return {"error": str(exc)}
        except Exception as exc:
            logger.exception("DocxEditorTool unexpected error: %s", exc)
            return {"error": f"Gagal mengedit file .docx: {exc}"}
